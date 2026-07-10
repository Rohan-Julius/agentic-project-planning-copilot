"""Day 4: document parsing tests (spec §4/§25) — PDF, DOCX, TXT, Markdown + negatives."""
from __future__ import annotations

import docx
import fitz
import pytest

from app.services.document_service import (
    CorruptedDocumentError,
    EmptyDocumentError,
    UnsupportedFormatError,
    parse_document,
)


def _make_pdf(tmp_path, pages: list[list[tuple[str, float]]]):
    """Each page is a list of (text, fontsize) lines, spaced well apart so PyMuPDF keeps
    them as distinct text blocks (verified empirically — insert_text() doesn't merge
    same-call lines into one block).
    """
    pdf = fitz.open()
    for lines in pages:
        page = pdf.new_page()
        y = 72.0
        for text, size in lines:
            page.insert_text((72, y), text, fontsize=size)
            y += size + 20
    path = tmp_path / "sample.pdf"
    pdf.save(path)
    pdf.close()
    return path


# --- TXT ---


def test_parse_txt_single_block(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("Users can log in with email and password.")

    parsed = parse_document(path, ".txt", "doc_1")

    assert parsed.document_id == "doc_1"
    assert len(parsed.blocks) == 1
    block = parsed.blocks[0]
    assert block.text == "Users can log in with email and password."
    assert block.page_number is None
    assert block.heading_level is None
    assert block.section_hierarchy_path == ""


def test_parse_empty_txt_raises(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n  ")

    with pytest.raises(EmptyDocumentError):
        parse_document(path, ".txt", "doc_1")


# --- Markdown ---


def test_parse_markdown_heading_hierarchy(tmp_path):
    path = tmp_path / "spec.md"
    path.write_text(
        "# Title\n"
        "\n"
        "Intro paragraph.\n"
        "\n"
        "## Section A\n"
        "\n"
        "Body under A.\n"
        "\n"
        "### Sub A.1\n"
        "\n"
        "Deep text under sub-section.\n"
    )

    parsed = parse_document(path, ".md", "doc_2")
    by_text = {b.text: b for b in parsed.blocks}

    assert by_text["Title"].heading_level == 1
    assert by_text["Title"].section_hierarchy_path == "Title"
    assert by_text["Intro paragraph."].section_hierarchy_path == "Title"

    assert by_text["Section A"].heading_level == 2
    assert by_text["Section A"].section_hierarchy_path == "Title > Section A"
    assert by_text["Body under A."].section_hierarchy_path == "Title > Section A"

    assert by_text["Sub A.1"].heading_level == 3
    assert by_text["Sub A.1"].section_hierarchy_path == "Title > Section A > Sub A.1"
    assert (
        by_text["Deep text under sub-section."].section_hierarchy_path
        == "Title > Section A > Sub A.1"
    )
    assert all(b.page_number is None for b in parsed.blocks)


def test_parse_markdown_sibling_resets_hierarchy(tmp_path):
    path = tmp_path / "spec.md"
    path.write_text("# Title\n\n## Section A\n\nBody A.\n\n## Section B\n\nBody B.\n")

    parsed = parse_document(path, ".md", "doc_3")
    by_text = {b.text: b for b in parsed.blocks}

    assert by_text["Section B"].section_hierarchy_path == "Title > Section B"
    assert by_text["Body B."].section_hierarchy_path == "Title > Section B"


# --- DOCX ---


def test_parse_docx_heading_hierarchy(tmp_path):
    document = docx.Document()
    document.add_heading("Title", level=1)
    document.add_paragraph("Intro paragraph.")
    document.add_heading("Section A", level=2)
    document.add_paragraph("Body under A.")
    path = tmp_path / "spec.docx"
    document.save(path)

    parsed = parse_document(path, ".docx", "doc_4")
    by_text = {b.text: b for b in parsed.blocks}

    assert by_text["Title"].heading_level == 1
    assert by_text["Section A"].heading_level == 2
    assert by_text["Section A"].section_hierarchy_path == "Title > Section A"
    assert by_text["Body under A."].section_hierarchy_path == "Title > Section A"
    assert all(b.page_number is None for b in parsed.blocks)


def test_parse_corrupted_docx_raises(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is not a real docx zip")

    with pytest.raises(CorruptedDocumentError):
        parse_document(path, ".docx", "doc_5")


# --- PDF ---


def test_parse_pdf_pages_and_font_size_headings(tmp_path):
    path = _make_pdf(
        tmp_path,
        pages=[
            [
                ("Chapter 1", 20),
                ("This is the first line of body text.", 11),
                ("Section 1.1", 16),
                ("Body text under section 1.1.", 11),
            ],
            [
                ("Body text continues on page two.", 11),
            ],
        ],
    )

    parsed = parse_document(path, ".pdf", "doc_6")
    by_text = {b.text: b for b in parsed.blocks}

    assert by_text["Chapter 1"].heading_level == 1
    assert by_text["Chapter 1"].page_number == 1
    assert by_text["Section 1.1"].heading_level == 2
    assert by_text["Section 1.1"].section_hierarchy_path == "Chapter 1 > Section 1.1"

    body_line = by_text["This is the first line of body text."]
    assert body_line.heading_level is None
    assert body_line.page_number == 1
    assert body_line.section_hierarchy_path == "Chapter 1"

    under_section = by_text["Body text under section 1.1."]
    assert under_section.section_hierarchy_path == "Chapter 1 > Section 1.1"

    page_two_line = by_text["Body text continues on page two."]
    assert page_two_line.page_number == 2
    # Heading hierarchy carries across pages until a new heading is seen.
    assert page_two_line.section_hierarchy_path == "Chapter 1 > Section 1.1"


def test_parse_empty_pdf_raises(tmp_path):
    pdf = fitz.open()
    pdf.new_page()
    path = tmp_path / "blank.pdf"
    pdf.save(path)
    pdf.close()

    with pytest.raises(EmptyDocumentError):
        parse_document(path, ".pdf", "doc_7")


def test_parse_corrupted_pdf_raises(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a real pdf file, just garbage bytes")

    with pytest.raises(CorruptedDocumentError):
        parse_document(path, ".pdf", "doc_8")


# --- Negative: unsupported format ---


def test_parse_unsupported_format_raises(tmp_path):
    path = tmp_path / "malware.exe"
    path.write_bytes(b"binary")

    with pytest.raises(UnsupportedFormatError):
        parse_document(path, ".exe", "doc_9")
