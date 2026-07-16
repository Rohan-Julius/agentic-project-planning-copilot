"""Document upload, storage, metadata/versioning, and parsing (spec §16.3, §20.2, §22, §4).

Files are written under `settings.documents_dir/<project_id>/` only — never outside the
project data directory (§20.2 filesystem restriction).

Parsing (Day 4, DESIGN.md §4) turns a stored file into a `ParsedDocument`: a flat list of
`ParsedBlock`s carrying text, page number (PDF only), heading level, and a running
section-hierarchy path — the normalized intermediate the Day-5 chunker consumes.
"""
from __future__ import annotations

import re
import uuid
from collections import Counter
from pathlib import Path

import docx
import fitz
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.config import Settings
from app.models.document import DocumentChunkMeta, DocumentRecord
from app.schemas.document import SUPPORTED_EXTENSIONS, ChunkPayload, ParsedBlock, ParsedDocument


class DocumentParsingError(ValueError):
    """Base for document-parsing failures (spec §25) — the API maps these to 422."""


class UnsupportedFormatError(DocumentParsingError):
    """Raised when an uploaded document's extension isn't PDF/DOCX/TXT/Markdown."""


class EmptyDocumentError(DocumentParsingError):
    """Raised when a document has zero extractable characters (spec §25)."""


class CorruptedDocumentError(DocumentParsingError):
    """Raised when a PDF/DOCX file can't be read by its parsing library (spec §25)."""


def generate_document_id() -> str:
    return f"doc_{uuid.uuid4().hex[:12]}"


def _project_dir(settings: Settings, project_id: str) -> Path:
    project_dir = settings.documents_dir / project_id
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir


def _next_version(session: Session, project_id: str, document_name: str) -> str:
    count = session.scalar(
        select(func.count())
        .select_from(DocumentRecord)
        .where(
            DocumentRecord.project_id == project_id,
            DocumentRecord.document_name == document_name,
        )
    )
    return f"{(count or 0) + 1}.0"


def save_uploaded_document(
    session: Session,
    settings: Settings,
    project_id: str,
    filename: str,
    content: bytes,
    document_type: str = "",
) -> DocumentRecord:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported document format '{extension}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    document_id = generate_document_id()
    dest = _project_dir(settings, project_id) / f"{document_id}{extension}"
    dest.write_bytes(content)

    document = DocumentRecord(
        document_id=document_id,
        project_id=project_id,
        document_name=filename,
        document_type=document_type,
        source_type="project",
        file_path=str(dest),
        document_version=_next_version(session, project_id, filename),
    )
    session.add(document)
    session.commit()
    session.refresh(document)
    return document


def save_text_document(
    session: Session,
    settings: Settings,
    project_id: str,
    document_name: str,
    text_content: str,
    document_type: str = "",
) -> DocumentRecord:
    filename = document_name if Path(document_name).suffix else f"{document_name}.txt"
    return save_uploaded_document(
        session,
        settings,
        project_id,
        filename,
        text_content.encode("utf-8"),
        document_type,
    )


def list_documents(session: Session, project_id: str) -> list[DocumentRecord]:
    return list(
        session.scalars(
            select(DocumentRecord).where(DocumentRecord.project_id == project_id)
        )
    )


def get_document(
    session: Session, project_id: str, document_id: str
) -> DocumentRecord | None:
    return session.scalar(
        select(DocumentRecord).where(
            DocumentRecord.project_id == project_id,
            DocumentRecord.document_id == document_id,
        )
    )


def delete_document(session: Session, project_id: str, document_id: str) -> bool:
    document = get_document(session, project_id, document_id)
    if document is None:
        return False
    file_path = Path(document.file_path)
    if file_path.exists():
        file_path.unlink()
    session.delete(document)
    session.commit()
    return True


# --- Parsing (Day 4, DESIGN.md §4) ---------------------------------------------------

_RawBlock = tuple[str, "int | None", "int | None"]  # (text, page_number, heading_level)

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_DOCX_HEADING_STYLE_RE = re.compile(r"^Heading (\d+)$", re.IGNORECASE)

_BOLD_FLAG = 1 << 4  # PyMuPDF span.flags bit for bold (§4 heuristic: font-size/bold spans)
_BOLD_FALLBACK_MAX_WORDS = 10


def _build_hierarchy_blocks(raw_blocks: list[_RawBlock]) -> list[ParsedBlock]:
    """Turn (text, page_number, heading_level) tuples into `ParsedBlock`s with a running
    section-hierarchy path (DESIGN.md §4): on a heading of level L, truncate the stack to
    L-1 and push the heading; body blocks inherit the current stack.
    """
    stack: list[str] = []
    blocks: list[ParsedBlock] = []
    for text, page_number, heading_level in raw_blocks:
        if heading_level is not None:
            stack[heading_level - 1 :] = [text]
        blocks.append(
            ParsedBlock(
                text=text,
                page_number=page_number,
                heading_level=heading_level,
                section_hierarchy_path=" > ".join(stack),
            )
        )
    return blocks


def _pdf_block_text_and_font(block: dict) -> tuple[str, float, bool]:
    """Concatenate a PyMuPDF text block's spans; report its max font size and whether
    every span in it is bold (both feed the heading heuristic below).
    """
    line_texts: list[str] = []
    max_size = 0.0
    all_bold = True
    any_span = False
    for line in block.get("lines", []):
        span_texts = []
        for span in line.get("spans", []):
            any_span = True
            span_texts.append(span["text"])
            max_size = max(max_size, span["size"])
            if not (span["flags"] & _BOLD_FLAG):
                all_bold = False
        span_text = "".join(span_texts).strip()
        if span_text:
            line_texts.append(span_text)
    text = " ".join(line_texts).strip()
    return text, max_size, (all_bold if any_span else False)


def _parse_pdf(file_path: Path) -> list[_RawBlock]:
    """PyMuPDF extraction with page numbers and a font-size/bold heading heuristic (§4).

    Heading level is primarily the rank of a block's font size among the sizes larger than
    the document's body size (the most common size, by character count). A short, fully
    bold, non-sentence-ending block at body size is treated as a fallback heading one level
    below the largest size-based level — this catches bold-only headings when a document
    doesn't vary font size, at the cost of not distinguishing sub-levels among them.
    """
    try:
        pdf = fitz.open(file_path)
    except Exception as exc:
        raise CorruptedDocumentError(f"Could not open PDF '{file_path.name}': {exc}") from exc

    try:
        candidates: list[tuple[int, str, float, bool]] = []
        size_weight: Counter[float] = Counter()
        for page_index in range(pdf.page_count):
            try:
                page_dict = pdf.load_page(page_index).get_text("dict")
            except Exception as exc:
                raise CorruptedDocumentError(
                    f"Could not read page {page_index + 1} of '{file_path.name}': {exc}"
                ) from exc
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # skip images
                    continue
                text, max_size, is_bold = _pdf_block_text_and_font(block)
                if not text:
                    continue
                rounded_size = round(max_size, 1)
                candidates.append((page_index + 1, text, rounded_size, is_bold))
                size_weight[rounded_size] += len(text)
    finally:
        pdf.close()

    if not candidates:
        return []

    body_size = size_weight.most_common(1)[0][0]
    heading_sizes = sorted({size for *_, size, _ in candidates if size > body_size}, reverse=True)
    size_to_level = {size: level for level, size in enumerate(heading_sizes, start=1)}
    bold_fallback_level = len(heading_sizes) + 1

    raw: list[_RawBlock] = []
    for page_number, text, size, is_bold in candidates:
        level = size_to_level.get(size)
        if (
            level is None
            and is_bold
            and len(text.split()) <= _BOLD_FALLBACK_MAX_WORDS
            and not text.rstrip().endswith((".", "?", "!", ":"))
        ):
            level = bold_fallback_level
        raw.append((text, page_number, level))
    return raw


def _parse_docx(file_path: Path) -> list[_RawBlock]:
    """python-docx paragraphs; heading level from the `Heading N` style name (§4)."""
    try:
        document = docx.Document(file_path)
    except Exception as exc:
        raise CorruptedDocumentError(f"Could not open DOCX '{file_path.name}': {exc}") from exc

    raw: list[_RawBlock] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        match = _DOCX_HEADING_STYLE_RE.match(style_name.strip())
        level = int(match.group(1)) if match else None
        raw.append((text, None, level))
    return raw


def _parse_text(file_path: Path, extension: str) -> list[_RawBlock]:
    """TXT: whole file as one block, no headings. Markdown: `#`/`##`/`###+` → level 1/2/3,
    paragraphs (blank-line-separated) as body blocks (§4).
    """
    content = file_path.read_text(encoding="utf-8", errors="replace")

    if extension == ".txt":
        text = content.strip()
        return [(text, None, None)] if text else []

    raw: list[_RawBlock] = []
    paragraph_lines: list[str] = []

    def flush() -> None:
        if paragraph_lines:
            raw.append((" ".join(paragraph_lines).strip(), None, None))
            paragraph_lines.clear()

    for line in content.splitlines():
        stripped = line.strip()
        heading_match = _MD_HEADING_RE.match(stripped)
        if heading_match:
            flush()
            level = min(len(heading_match.group(1)), 3)
            raw.append((heading_match.group(2).strip(), None, level))
        elif not stripped:
            flush()
        else:
            paragraph_lines.append(stripped)
    flush()
    return raw


def parse_document(file_path: Path, extension: str, document_id: str) -> ParsedDocument:
    """Dispatch on extension (§4) and normalize into a `ParsedDocument`.

    Raises `UnsupportedFormatError` for extensions outside `SUPPORTED_EXTENSIONS`,
    `CorruptedDocumentError` if the underlying library can't read the file, and
    `EmptyDocumentError` if zero extractable characters remain (§25 negatives).
    """
    extension = extension.lower()
    if extension == ".pdf":
        raw_blocks = _parse_pdf(file_path)
    elif extension == ".docx":
        raw_blocks = _parse_docx(file_path)
    elif extension in {".txt", ".md"}:
        raw_blocks = _parse_text(file_path, extension)
    else:
        raise UnsupportedFormatError(
            f"Unsupported document format '{extension}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    blocks = _build_hierarchy_blocks(raw_blocks)
    if not any(block.text.strip() for block in blocks):
        raise EmptyDocumentError(f"Document '{file_path.name}' has no extractable text")
    return ParsedDocument(document_id=document_id, blocks=blocks)


# --- Chunk metadata mirror (Day 5, DESIGN.md §5.2/§12.2) ----------------------------


def save_chunk_meta(session: Session, chunks: list[ChunkPayload]) -> None:
    """Upsert §12.2 chunk metadata into `document_chunk_meta`, mirroring the Qdrant
    payload (minus vectors/text) so the Day-13 citation validator can check "does this
    chunk_id exist and is it non-deleted" with a fast indexed SQL query.
    """
    for chunk in chunks:
        row = session.scalar(
            select(DocumentChunkMeta).where(DocumentChunkMeta.chunk_id == chunk.chunk_id)
        )
        if row is None:
            row = DocumentChunkMeta(chunk_id=chunk.chunk_id)
            session.add(row)
        row.document_id = chunk.document_id
        row.project_id = chunk.project_id
        row.page_number = chunk.page_number
        row.section = chunk.section
        row.section_hierarchy_path = chunk.section_hierarchy_path
        row.document_version = chunk.document_version
    session.commit()
